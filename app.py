import streamlit as st
import json
import os
import numpy as np
from docx import Document
import pdfplumber
import matplotlib.colors as mcolors
from operator import index
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
import string
from dotenv import load_dotenv
import docx
from transformers import T5ForConditionalGeneration, T5Tokenizer
from PIL import Image
from openai import OpenAI


load_dotenv()


session_state = st.session_state
if "user_index" not in st.session_state:
    st.session_state["user_index"] = 0


def signup(json_file_path="data.json"):
    st.title("Signup Page")
    with st.form("signup_form"):
        st.write("Fill in the details below to create an account:")
        name = st.text_input("Name:")
        email = st.text_input("Email:")
        age = st.number_input("Age:", min_value=0, max_value=120)
        sex = st.radio("Sex:", ("Male", "Female", "Other"))
        password = st.text_input("Password:", type="password")
        confirm_password = st.text_input("Confirm Password:", type="password")

        if st.form_submit_button("Signup"):
            if password == confirm_password:
                user = create_account(name, email, age, sex, password, json_file_path)
                session_state["logged_in"] = True
                session_state["user_info"] = user
            else:
                st.error("Passwords do not match. Please try again.")
# Load the trained model
model = T5ForConditionalGeneration.from_pretrained("/checkpoints/Model")
# Load the tokenizer
tokenizer = T5Tokenizer.from_pretrained("t5-small")

# Function to tokenize and lemmatize text
def tokenize_and_lemmatize(text):
    tokens = word_tokenize(text.lower())
    lemmatizer = WordNetLemmatizer()
    lemmatized_tokens = [lemmatizer.lemmatize(token) for token in tokens]
    return lemmatized_tokens

def check_login(username, password, json_file_path="data.json"):
    try:
        with open(json_file_path, "r") as json_file:
            data = json.load(json_file)

        for user in data["users"]:
            if user["email"] == username and user["password"] == password:
                session_state["logged_in"] = True
                session_state["user_info"] = user
                st.success("Login successful!")
                render_dashboard(user)
                return user

        st.error("Invalid credentials. Please try again.")
        return None
    except Exception as e:
        st.error(f"Error checking login: {e}")
        return None
# Function to generate summaries
def generate_summary(input_text):
    inputs = tokenizer(input_text, return_tensors="pt", max_length=512, truncation=True)
    outputs = model.generate(input_ids=inputs["input_ids"], attention_mask=inputs["attention_mask"], max_length=150, num_beams=1, early_stopping=True)
    generated_summary = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return generated_summary

def initialize_database(json_file_path="data.json"):
    try:
        # Check if JSON file exists
        if not os.path.exists(json_file_path):
            # Create an empty JSON structure
            data = {"users": []}
            with open(json_file_path, "w") as json_file:
                json.dump(data, json_file)
    except Exception as e:
        print(f"Error initializing database: {e}")


def create_account(name, email, age, sex, password, json_file_path="data.json"):
    try:
        # Check if the JSON file exists or is empty
        if not os.path.exists(json_file_path) or os.stat(json_file_path).st_size == 0:
            data = {"users": []}
        else:
            with open(json_file_path, "r") as json_file:
                data = json.load(json_file)

        # Append new user data to the JSON structure
        user_info = {
            "name": name,
            "email": email,
            "age": age,
            "sex": sex,
            "password": password,
            "resume": None,
            "job_description": None,
            "job_applied": None,
            'score': '0',
            "questions": None,
        }
        data["users"].append(user_info)

        # Save the updated data to JSON
        with open(json_file_path, "w") as json_file:
            json.dump(data, json_file, indent=4)

        st.success("Account created successfully! You can now login.")
        return user_info
    except json.JSONDecodeError as e:
        st.error(f"Error decoding JSON: {e}")
        return None
    except Exception as e:
        st.error(f"Error creating account: {e}")
        return None



def login(json_file_path="data.json"):
    st.title("Login Page")
    username = st.text_input("Username:")
    password = st.text_input("Password:", type="password")

    login_button = st.button("Login")

    if login_button:
        user = check_login(username, password, json_file_path)
        if user is not None:
            session_state["logged_in"] = True
            session_state["user_info"] = user
        else:
            st.error("Invalid credentials. Please try again.")

def get_user_info(email, json_file_path="data.json"):
    try:
        with open(json_file_path, "r") as json_file:
            data = json.load(json_file)
            for user in data["users"]:
                if user["email"] == email:
                    return user
        return None
    except Exception as e:
        st.error(f"Error getting user information: {e}")
        return None


def render_dashboard(user_info, json_file_path="data.json"):
    try:
        st.title(f"Welcome to the Dashboard, {user_info['name']}!")
        st.subheader("User Information:")
        st.write(f"Name: {user_info['name']}")
        st.write(f"Sex: {user_info['sex']}")
        st.write(f"Age: {user_info['age']}")
    except Exception as e:
        st.error(f"Error rendering dashboard: {e}")


def extract_text(file) -> str:
    if isinstance(file, str):
        file_extension = os.path.splitext(file)[1].lower()
    else:
        file_extension = os.path.splitext(file.name)[1].lower()
    
    if file_extension == '.pdf':
        if isinstance(file, str):
            with pdfplumber.open(file) as pdf:
                text = '\n'.join(page.extract_text() for page in pdf.pages if page.extract_text())
        else:
            with pdfplumber.open(file) as pdf:
                text = '\n'.join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif file_extension == '.docx':
        if isinstance(file, str):
            doc = docx.Document(file)
        else:
            doc = docx.Document(file)
        text = '\n'.join([para.text for para in doc.paragraphs])
    else:
        if isinstance(file, str):
            with open(file, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        else:
            with file as f:
                text = f.read()
    return text

def process_text(text):
    tokens = word_tokenize(text)
    tokens = [word.lower() for word in tokens if word.isalpha()]
    stop_words = set(stopwords.words('english'))
    tokens = [word for word in tokens if word not in stop_words]
    return tokens

def calculate_score(job_description_tokens, resume_tokens):
    job_description_freq = FreqDist(job_description_tokens)
    total_tokens_in_job_description = len(job_description_tokens)
    resume_tokens = list(set(resume_tokens))
    score = sum(job_description_freq[token] for token in resume_tokens)
    score_percentage = (score / total_tokens_in_job_description) * 100
    return score_percentage

def extract_keywords_from_resume(resume_text):
    resume_text = resume_text.lower()
    resume_tokens = word_tokenize(resume_text)
    resume_tokens = [
        token for token in resume_tokens if token not in string.punctuation
    ]
    stop_words = set(stopwords.words("english"))
    resume_tokens = [token for token in resume_tokens if token not in stop_words]
    processed_resume_text = " ".join(resume_tokens)
    client = OpenAI(
        api_key=os.environ.get("OPENAI_API_KEY"),
    )
    prompt = f"Extract top most important skill keywords from the given resume text:\n{processed_resume_text}\nKeywords:"
    messages = [{"role": "system", "content": prompt}]
    response = client.chat.completions.create(
        messages=messages,
        model="gpt-3.5-turbo",
    )
    return response.choices[0].message.content.split(",")

def generate_question(resume_text, job_description_text, candidate_name, previous_response=None, previous_question=None):
    prompt = f"Resume Text: {resume_text}\nJob Description: {job_description_text}\nCandidate Name: {candidate_name}\n"
    if previous_response and previous_question:
        prompt += f"Previous Response: {previous_response}\nPrevious Question: {previous_question}\n"
    client = OpenAI(
        api_key=os.environ.get("OPENAI_API_KEY"),
    )
    messages=[
            {"role": "system", "content": "You are the interviewer."},
            {"role": "system", "content": "You are interviewing a candidate. Ask a question based on the resume and job description. If the candidate has already answered a question, you can ask a follow-up question based on their response."},
            {"role": "user", "content": prompt}
        ]
    response = client.chat.completions.create(
        messages=messages,
        model="gpt-3.5-turbo",
    )
    return response.choices[0].message.content


def main(json_file_path="data.json"):
    st.sidebar.title("Resume Screening system")
    page = st.sidebar.radio(
        "Go to",
        ("Signup/Login", "Upload Resume", "Resume Analysis", "Generate Questions", "Evaluate Scores"),
        key="GET YOUR RESUME ANALYZED AND COMPARED",
    )

    if page == "Signup/Login":
        st.title("Signup/Login Page")
        login_or_signup = st.radio(
            "Select an option", ("Login", "Signup"), key="login_signup"
        )
        if login_or_signup == "Login":
            login(json_file_path)
        else:
            signup(json_file_path)

    elif page == "Upload Resume":
        if session_state.get("logged_in"):
            uploaded_file = st.file_uploader("Choose a file", type=None)
            if uploaded_file is not None:
                resume_text = extract_text(uploaded_file)
                st.write("File name: ", uploaded_file.name)
                st.success("File uploaded successfully!")
                st.image(Image.open('Images/logo.png'), use_column_width=True)
                with open(json_file_path, "r+") as json_file:
                    data = json.load(json_file)
                    user_index = next((i for i, user in enumerate(data["users"]) if user["email"] == session_state["user_info"]["email"]), None)
                    if user_index is not None:
                        user_info = data["users"][user_index]
                        user_info["resume"] = resume_text
                        session_state["user_info"] = user_info
                        json_file.seek(0)
                        json.dump(data, json_file, indent=4)
                        json_file.truncate()
                    else:
                        st.error("User not found.")
                        
            st.subheader("Select a role to you want to apply for:")
            
            BASE_DIR = "Data\\JobDesc\\"
            job_description = st.selectbox(
                "Select a role",
                [
                    "-Select-",
                    "Backend Developer",
                    "Billing cum Logistics Manager",
                    "Data Scientist",
                    "Director of Engineering",
                    "Global Industry Content Manager",
                    "HTML Developer",
                    "IT Project Manager",
                    "Lead Technical Program Manager",
                    "Primary English Teacher",
                    "Revenue Reporting Data Analyst",
                    "Senior Product Manager",
                    "Senior Software Developer",
                    "Web Developer",
                    "Web_dev_job",
                ],
                key="job_description",
            )
            if job_description and job_description != "-Select-":
                file_path = os.path.join(BASE_DIR, f"{job_description}.docx")
                job_description_text = extract_text(file_path)
                st.subheader("Job Description:")
                st.write(job_description_text)
                if st.button("Apply"):
                    with open(json_file_path, "r+") as json_file:
                        data = json.load(json_file)
                        user_index = next((i for i, user in enumerate(data["users"]) if user["email"] == session_state["user_info"]["email"]), None)
                        if user_index is not None:
                            user_info = data["users"][user_index]
                            user_info["job_description"] = job_description_text
                            user_info["job_applied"] = job_description
                            session_state["user_info"] = user_info
                            json_file.seek(0)
                            json.dump(data, json_file, indent=4)
                            json_file.truncate()
                        else:
                            st.error("User not found.")
                    st.success("Job application submitted successfully!")
                    
        else:
            st.warning("Please login/signup to view the dashboard.")

    elif page == "Resume Analysis":
        if session_state.get("logged_in"):
            st.title("Get Your Resume Analyzed and Compared")
            resume_text = session_state["user_info"]["resume"]
            resume_keywords = extract_keywords_from_resume(resume_text)
            print(resume_keywords)
            st.subheader("Skills of the candidate:")
            for keyword in resume_keywords:
                st.write(f"- {keyword.strip()}")
            job_description_text = session_state["user_info"]["job_description"]
            job_description_tokens = process_text(job_description_text)
            resume_tokens = process_text(resume_text)
            score = calculate_score(job_description_tokens, resume_tokens)
            with open(json_file_path, "r+") as json_file:
                data = json.load(json_file)
                user_index = next((i for i, user in enumerate(data["users"]) if user["email"] == session_state["user_info"]["email"]), None)
                if user_index is not None:
                    user_info = data["users"][user_index]
                    user_info["score"] = score
                    session_state["user_info"] = user_info
                    json_file.seek(0)
                    json.dump(data, json_file, indent=4)
                    json_file.truncate()
                else:
                    st.error("User not found.")
            # Visualization
            st.header("Resume Score")
            if score >= 60:
                st.success(f"Congratulations! Your resume matches {score}% with the job description.")
            elif score >= 20:
                st.warning(f"Your resume matches {score}% with the job description. Consider improving it for better results.")
            else:
                st.error(f"Your resume matches only {score}% with the job description. Consider significant improvements.")
            percentage_score = score / 100
            percentage_remainder = 1 - percentage_score

            # Create a Plotly figure for the pie chart
            fig = go.Figure(data=[go.Pie(labels=['Matched', 'Unmatched'], 
                                        values=[percentage_score, percentage_remainder], 
                                        hole=0.3,
                                        marker_colors=['rgba(0, 128, 0, 0.7)', 'rgba(255, 0, 0, 0.7)'])])

            # Update layout to add title
            fig.update_layout(title_text="Resume Score")

            # Display the chart
            st.plotly_chart(fig)
            # compare with other candidates
            st.subheader("How does your resume compare with other candidates?")
            role = session_state["user_info"]["job_applied"]
            scores = [int(user["score"]) for user in data["users"] if user["job_applied"] == role]
            # Plot an interactive graph
            fig = go.Figure()
            fig.add_trace(go.Histogram(x=scores, 
                                        histnorm='percent',
                                        marker_color='rgba(0, 0, 255, 0.7)',
                                        opacity=0.75))

            fig.update_layout(title_text=f'Distribution of Scores for Job Role: {role}',
                            xaxis_title='Resume Score',
                            yaxis_title='Percentage of Candidates',
                            bargap=0.05)

            st.plotly_chart(fig)
            
    elif page == "Generate Questions":
        if session_state.get("logged_in"):
            user_info = session_state["user_info"]
            st.title("Interview Question Generator Chat")
            st.subheader("Generate Interview Questions")
            if session_state["user_info"]["resume"] is None or session_state["user_info"]["job_description"] is None:
                st.warning("Please upload your resume and apply for a job to generate interview questions.")
                return
                
            st.markdown("### Interview Questions")
            if user_info["questions"] is None:
                previous_response = None
                previous_question = None
            else:
                previous_response = user_info["questions"][-1]["response"]
                previous_question = user_info["questions"][-1]["question"]
            
            question = generate_question(session_state["user_info"]["resume"], session_state["user_info"]["job_description"], session_state["user_info"]["name"], previous_response, previous_question)
            st.markdown(question)
            response = st.text_input("Your response:")
            if st.button("Next Question"):
                if response is not None and len(response) > 0:
                    with open(json_file_path, "r+") as json_file:
                        data = json.load(json_file)
                        user_index = next((i for i, user in enumerate(data["users"]) if user["email"] == session_state["user_info"]["email"]), None)
                        if user_index is not None:
                            user_info = data["users"][user_index]
                            if user_info["questions"] is None:
                                user_info["questions"] = []
                            user_info["questions"].append({"question": question, "response": response})
                            session_state["user_info"] = user_info
                            json_file.seek(0)
                            json.dump(data, json_file, indent=4)
                            json_file.truncate()
                        else:
                            st.error("User not found.")
                    st.rerun()
            if st.button("Finish"):
                with open(json_file_path, "r+") as json_file:
                    data = json.load(json_file)
                    user_index = next((i for i, user in enumerate(data["users"]) if user["email"] == session_state["user_info"]["email"]), None)
                    if user_index is not None:
                        user_info = data["users"][user_index]
                        user_info["questions"].append({"question": question, "response": response})
                        session_state["user_info"] = user_info
                        json_file.seek(0)
                        json.dump(data, json_file, indent=4)
                        json_file.truncate()
                    else:
                        st.error("User not found.")
                st.success("Interview questions completed successfully!")
                return
            
        else:
            st.warning("Please login/signup to view the dashboard.")
    elif page == "Evaluate Scores":
        if session_state.get("logged_in"):
            user_info = session_state["user_info"]
            st.title("Evaluate Scores")
            st.subheader("Evaluate Scores")
            if user_info["questions"] is not None:
                st.markdown("### Interview Questions and Responses")
                for question in user_info["questions"]:
                    st.markdown(f"**Question:** {question['question']}")
                    st.markdown(f"**Response:** {question['response']}")
                    st.markdown("---")
                st.subheader("Score")
                
                # score is the mean of the scores of the responses to the questions
                score = 0
                count = 0
                for question in user_info["questions"]:
                    count += 1
                    response = question["response"]
                    ques = question["question"]
                    score += calculate_score(process_text(ques), process_text(response))
                score = round(score/count, 2)

                st.write(f"Score: {score} %")
                
        else:
            st.warning("Please login/signup to view the dashboard.")
            
            
                
if __name__ == "__main__":
    import os
    import nltk
    initialize_database()
    nltk.download('punkt')
    nltk.download('stopwords')
    main()
